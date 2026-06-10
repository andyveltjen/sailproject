#!/usr/bin/env python3
"""
docx2wiki.py — Converteer een Word-document naar een Grav wiki pagina.

Gebruik:
    python3 docx2wiki.py bestand.docx pad/naar/wiki/pagina

Voorbeeld:
    python3 docx2wiki.py mijnpagina.docx 02.wiki/02.literatuuronderzoek/03.nieuwe-pagina

Het script maakt de map aan (als die nog niet bestaat) en schrijft docs.nl.md.
"""

import sys
import os
import re
from docx import Document
from docx.oxml.ns import qn


def get_paragraph_style(para):
    """Geef de stijlnaam terug, genormaliseerd naar lowercase."""
    return para.style.name.lower().strip()


def get_inline_markup(para):
    """Zet de runs van een paragraaf om naar Markdown met bold/italic."""
    result = ""
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        result += text
    return result.strip()


def get_list_level(para):
    """Geef het inspring-niveau van een lijstparagraaf terug (0-based)."""
    numPr = para._element.find(qn("w:pPr") + "/" + qn("w:numPr"))
    if numPr is None:
        numPr = para._element.find(".//{%s}numPr" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
    if numPr is not None:
        ilvl = numPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
        if ilvl is not None:
            return int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", 0))
    return 0


def get_num_info(para, doc):
    """
    Geef (is_numbered, level) terug voor een lijstparagraaf.
    Kijkt in de numbering-definitie om bullet vs. decimal te onderscheiden.
    """
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    numPr = para._element.find(f".//{{{ns}}}numPr")
    if numPr is None:
        return None, 0

    ilvl_el = numPr.find(f"{{{ns}}}ilvl")
    numId_el = numPr.find(f"{{{ns}}}numId")
    if ilvl_el is None or numId_el is None:
        return None, 0

    ilvl = int(ilvl_el.get(f"{{{ns}}}val", 0))
    numId = numId_el.get(f"{{{ns}}}val")

    # Zoek het numFmt op in de numbering-part
    try:
        numbering = doc.part.numbering_part
        if numbering is None:
            return "bullet", ilvl
        root = numbering._element
        # Zoek de abstractNumId via numId
        abstract_id = None
        for num in root.findall(f"{{{ns}}}num"):
            if num.get(f"{{{ns}}}numId") == numId:
                ref = num.find(f"{{{ns}}}abstractNumId")
                if ref is not None:
                    abstract_id = ref.get(f"{{{ns}}}val")
                break
        if abstract_id is None:
            return "bullet", ilvl
        # Zoek het level en numFmt
        for absNum in root.findall(f"{{{ns}}}abstractNum"):
            if absNum.get(f"{{{ns}}}abstractNumId") == abstract_id:
                for lvl in absNum.findall(f"{{{ns}}}lvl"):
                    if lvl.get(f"{{{ns}}}ilvl") == str(ilvl):
                        fmt_el = lvl.find(f"{{{ns}}}numFmt")
                        fmt = fmt_el.get(f"{{{ns}}}val") if fmt_el is not None else "bullet"
                        return fmt, ilvl
    except Exception:
        pass

    return "bullet", ilvl


def is_numbered_list(para):
    """Controleer of dit een genummerde lijst is (numFmt decimal)."""
    numPr = para._element.find(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
    )
    return numPr is not None


def extract_title_from_doc(doc):
    """Zoek de eerste Title-stijl paragraaf en gebruik die als paginatitel."""
    for para in doc.paragraphs:
        style = get_paragraph_style(para)
        if style == "title" and para.text.strip():
            return para.text.strip()
    # Fallback: eerste heading 1
    for para in doc.paragraphs:
        style = get_paragraph_style(para)
        if ("heading 1" in style or "kop 1" in style) and para.text.strip():
            return para.text.strip()
    return "Nieuwe pagina"


def convert_table(table):
    """Zet een Word-tabel om naar Markdown-tabel."""
    lines = []
    rows = table.rows
    if not rows:
        return ""

    for i, row in enumerate(rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Scheidingslijn na de header
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(lines)


def convert_doc_to_markdown(doc_path):
    """Hoofdfunctie: verwerk het .docx bestand naar Markdown-inhoud."""
    doc = Document(doc_path)
    title = extract_title_from_doc(doc)

    # Bouw de lijst van block-elementen op (paragrafen + tabellen in volgorde)
    # We werken met doc.element.body zodat tabellen en paragrafen in volgorde komen
    blocks = []
    body = doc.element.body
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    para_index = 0
    table_index = 0
    all_paras = doc.paragraphs
    all_tables = doc.tables

    para_elements = [p._element for p in all_paras]
    table_elements = [t._element for t in all_tables]

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_index < len(all_paras):
                blocks.append(("para", all_paras[para_index]))
                para_index += 1
        elif tag == "tbl":
            if table_index < len(all_tables):
                blocks.append(("table", all_tables[table_index]))
                table_index += 1

    # Nummering bijhouden voor genummerde lijsten per niveau
    counters = [0] * 10
    prev_was_list = False

    lines = []
    skip_title = True  # sla de eerste Title-paragraaf over (wordt frontmatter)

    for block_type, block in blocks:
        if block_type == "table":
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(convert_table(block))
            lines.append("")
            prev_was_list = False
            continue

        para = block
        style = get_paragraph_style(para)
        text = get_inline_markup(para)
        raw_text = para.text.strip()

        if not raw_text:
            if prev_was_list:
                pass  # geen extra lege regel midden in lijst
            else:
                lines.append("")
            prev_was_list = False
            continue

        # Title → alleen de eerste overslaan (wordt frontmatter-titel)
        if style == "title":
            if skip_title:
                skip_title = False
                prev_was_list = False
                continue
            else:
                lines.append(f"# {text}")
                lines.append("")
                prev_was_list = False
                continue

        # Heading 1 / Kop 1
        if "heading 1" in style or style in ("kop 1", "heading1"):
            counters = [0] * 10  # reset alle lijsttellers bij nieuwe sectie
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(f"## {text}")
            lines.append("")
            prev_was_list = False
            continue

        # Heading 2 / Kop 2
        if "heading 2" in style or style in ("kop 2", "heading2"):
            counters[1:] = [0] * (len(counters) - 1)  # reset sub-niveaus
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(f"### {text}")
            lines.append("")
            prev_was_list = False
            continue

        # Heading 3 / Kop 3
        if "heading 3" in style or style in ("kop 3", "heading3"):
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(f"#### {text}")
            lines.append("")
            prev_was_list = False
            continue

        # Genummerde lijst — detecteer op stijlnaam (betrouwbaarder dan numPr)
        # "List Number" = niveau 1, "List Number 2" = niveau 2, enz.
        list_num_match = re.match(
            r"^(list number|lijstnummer|genummerde lijst)( (\d+))?$", style
        )
        if list_num_match:
            level = int(list_num_match.group(3) or 1) - 1  # 0-based
            for i in range(level + 1, len(counters)):
                counters[i] = 0
            counters[level] += 1
            indent = "   " * level
            lines.append(f"{indent}{counters[level]}. {text}")
            prev_was_list = True
            continue

        # Bullet lijst (List Bullet stijlen)
        if re.match(r"^(list bullet|lijst met opsommingstekens|lijstopsomming)", style):
            level = get_list_level(para)
            indent = "   " * level
            lines.append(f"{indent}- {text}")
            prev_was_list = True
            continue

        # "List Paragraph" — generieke lijststijl: kijk op numFmt in numbering-definitie
        if style == "list paragraph":
            fmt, level = get_num_info(para, doc)
            if fmt == "decimal":
                for i in range(level + 1, len(counters)):
                    counters[i] = 0
                counters[level] += 1
                indent = "   " * level
                lines.append(f"{indent}{counters[level]}. {text}")
            else:
                indent = "   " * level
                lines.append(f"{indent}- {text}")
            prev_was_list = True
            continue

        # Blockquote / citaat
        if "quote" in style or "citaat" in style or "block text" in style:
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(f"> {text}")
            lines.append("")
            prev_was_list = False
            continue

        # Normale paragraaf
        if lines and lines[-1] != "" and prev_was_list:
            lines.append("")
        lines.append(text)
        lines.append("")
        prev_was_list = False

    # Verwijder overbodige lege regels aan het einde
    while lines and lines[-1] == "":
        lines.pop()

    # Meerdere opeenvolgende lege regels samenvoegen tot één
    cleaned = []
    prev_empty = False
    for line in lines:
        if line == "":
            if not prev_empty:
                cleaned.append(line)
            prev_empty = True
        else:
            cleaned.append(line)
            prev_empty = False

    return title, "\n".join(cleaned)


def build_frontmatter(title):
    return f"""---
title: '{title}'
taxonomy:
    category:
        - docs
---

"""


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    docx_path = sys.argv[1]
    wiki_path = sys.argv[2]

    if not os.path.isfile(docx_path):
        print(f"Fout: bestand niet gevonden: {docx_path}")
        sys.exit(1)

    # Bepaal de uitvoermap relatief aan de wiki-pages map
    base = "/Applications/MAMP/htdocs/sailproject/user/pages"
    out_dir = os.path.join(base, wiki_path)
    out_file = os.path.join(out_dir, "docs.nl.md")

    print(f"Verwerk: {docx_path}")
    title, body = convert_doc_to_markdown(docx_path)
    print(f"Titel gevonden: {title}")

    os.makedirs(out_dir, exist_ok=True)

    content = build_frontmatter(title) + body + "\n"

    with open(out_file, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"Geschreven naar: {out_file}")
    print()
    print("Volgende stap: controleer het bestand en commit met git.")


if __name__ == "__main__":
    main()
